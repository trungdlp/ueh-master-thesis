"""Build UEH-compliant reference.docx for pandoc.

Usage: python build_reference.py [output_path]
"""
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---- Page setup: A4, UEH margins ----
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(3.0)
section.bottom_margin = Cm(3.0)
section.left_margin = Cm(3.5)
section.right_margin = Cm(2.0)
section.header_distance = Cm(1.5)
section.footer_distance = Cm(1.5)

FONT = "Times New Roman"

def set_font(run_or_rpr, name=FONT, size=13, bold=False, italic=False):
    rpr = run_or_rpr
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), name)
    # size
    for tag in ("w:sz", "w:szCs"):
        e = rpr.find(qn(tag))
        if e is None:
            e = OxmlElement(tag)
            rpr.append(e)
        e.set(qn("w:val"), str(size * 2))
    if bold:
        b = OxmlElement("w:b")
        rpr.append(b)
        bcs = OxmlElement("w:bCs")
        rpr.append(bcs)
    if italic:
        i = OxmlElement("w:i")
        rpr.append(i)
    # Force black color
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    rpr.append(color)

def style_paragraph(style, *, size=13, bold=False, italic=False,
                    align=None, line_spacing=1.5, first_indent=None,
                    space_before=0, space_after=0, keep_with_next=False):
    # paragraph properties
    pf = style.paragraph_format
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_indent is not None:
        pf.first_line_indent = first_indent
    if align is not None:
        style.paragraph_format.alignment = align
    if keep_with_next:
        pPr = style.element.get_or_add_pPr()
        kwn = OxmlElement("w:keepNext")
        pPr.append(kwn)

    # run/font properties
    rpr = style.element.get_or_add_rPr()
    # clear existing font tags
    for tag in ("w:rFonts", "w:sz", "w:szCs", "w:b", "w:bCs", "w:i", "w:iCs", "w:color"):
        for e in rpr.findall(qn(tag)):
            rpr.remove(e)
    set_font(rpr, FONT, size, bold, italic)

styles = doc.styles

# Normal (body)
normal = styles["Normal"]
style_paragraph(
    normal, size=13, bold=False,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    line_spacing=1.5,
    first_indent=Cm(1.0),
)

# Heading 1 -> Chương: 14pt bold UPPERCASE centered
h1 = styles["Heading 1"]
style_paragraph(
    h1, size=14, bold=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    line_spacing=1.5, first_indent=Cm(0),
    space_before=12, space_after=12, keep_with_next=True,
)
# add caps
h1_rpr = h1.element.get_or_add_rPr()
caps = OxmlElement("w:caps")
h1_rpr.append(caps)

# Heading 2: 13pt bold left
h2 = styles["Heading 2"]
style_paragraph(
    h2, size=13, bold=True,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    line_spacing=1.5, first_indent=Cm(0),
    space_before=12, space_after=6, keep_with_next=True,
)

# Heading 3: 13pt bold italic
h3 = styles["Heading 3"]
style_paragraph(
    h3, size=13, bold=True, italic=True,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    line_spacing=1.5, first_indent=Cm(0),
    space_before=6, space_after=6, keep_with_next=True,
)

# Title (used by pandoc for document title) - reuse for cover
if "Title" in [s.name for s in styles]:
    title = styles["Title"]
    style_paragraph(
        title, size=16, bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.5, first_indent=Cm(0),
        space_before=12, space_after=12,
    )

# TOC heading style if exists
for name in ("TOC Heading", "toc 1", "toc 2", "toc 3"):
    try:
        s = styles[name]
        style_paragraph(
            s, size=13, bold=(name == "TOC Heading"),
            align=WD_ALIGN_PARAGRAPH.LEFT,
            line_spacing=1.5, first_indent=Cm(0),
        )
    except KeyError:
        pass

# Footer with page number, centered
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run()
# page number field
fldChar1 = OxmlElement("w:fldChar")
fldChar1.set(qn("w:fldCharType"), "begin")
instr = OxmlElement("w:instrText")
instr.text = "PAGE"
fldChar2 = OxmlElement("w:fldChar")
fldChar2.set(qn("w:fldCharType"), "end")
run._r.append(fldChar1)
run._r.append(instr)
run._r.append(fldChar2)
run.font.name = FONT
run.font.size = Pt(13)

# ---- Table style with borders + 100% width ----
# Pandoc applies style "Table" to its tables. We define/override that style here.
from docx.oxml import parse_xml

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
nsmap_decl = f'xmlns:w="{W_NS}"'

table_style_xml = f'''
<w:style {nsmap_decl} w:type="table" w:styleId="Table">
  <w:name w:val="Table"/>
  <w:basedOn w:val="TableNormal"/>
  <w:uiPriority w:val="59"/>
  <w:rPr>
    <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>
    <w:sz w:val="26"/>
    <w:szCs w:val="26"/>
    <w:color w:val="000000"/>
  </w:rPr>
  <w:tblPr>
    <w:tblW w:w="5000" w:type="pct"/>
    <w:jc w:val="center"/>
    <w:tblBorders>
      <w:top w:val="single" w:sz="6" w:space="0" w:color="000000"/>
      <w:left w:val="single" w:sz="6" w:space="0" w:color="000000"/>
      <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>
      <w:right w:val="single" w:sz="6" w:space="0" w:color="000000"/>
      <w:insideH w:val="single" w:sz="6" w:space="0" w:color="000000"/>
      <w:insideV w:val="single" w:sz="6" w:space="0" w:color="000000"/>
    </w:tblBorders>
    <w:tblCellMar>
      <w:top w:w="80" w:type="dxa"/>
      <w:left w:w="120" w:type="dxa"/>
      <w:bottom w:w="80" w:type="dxa"/>
      <w:right w:w="120" w:type="dxa"/>
    </w:tblCellMar>
  </w:tblPr>
  <w:tblStylePr w:type="firstRow">
    <w:rPr>
      <w:b/>
      <w:bCs/>
      <w:color w:val="000000"/>
    </w:rPr>
    <w:tblPr/>
    <w:trPr>
      <w:tblHeader/>
    </w:trPr>
    <w:tcPr>
      <w:shd w:val="clear" w:color="auto" w:fill="D9D9D9"/>
    </w:tcPr>
  </w:tblStylePr>
</w:style>
'''

styles_element = styles.element
# Remove any existing Table style to avoid duplicates
for existing in styles_element.findall(qn("w:style")):
    if existing.get(qn("w:styleId")) == "Table":
        styles_element.remove(existing)
styles_element.append(parse_xml(table_style_xml))

out_path = sys.argv[1] if len(sys.argv) > 1 else "reference-ueh.docx"
doc.save(out_path)
print(f"OK -> {out_path}")
